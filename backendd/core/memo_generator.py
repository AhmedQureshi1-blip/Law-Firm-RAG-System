from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

# ── Colour constants ───────────────────────────────────────────────────
COLOUR_PRIMARY   = RGBColor(0x14, 0x3C, 0x6E)   # dark blue
COLOUR_ACCENT    = RGBColor(0x00, 0x70, 0xA4)   # mid blue
COLOUR_RED       = RGBColor(0xAA, 0x23, 0x23)   # risk red
COLOUR_AMBER     = RGBColor(0xB8, 0x5C, 0x00)   # medium amber
COLOUR_GREEN     = RGBColor(0x1E, 0x6E, 0x37)   # low green
COLOUR_LIGHT     = RGBColor(0xF5, 0xF5, 0xF5)   # light gray


def _risk_colour(level: str) -> RGBColor:
    mapping = {
        "HIGH"  : COLOUR_RED,
        "MEDIUM": COLOUR_AMBER,
        "LOW"   : COLOUR_GREEN,
    }
    return mapping.get(level.upper(), COLOUR_ACCENT)


def _set_cell_bg(cell, hex_colour: str):
    """Set table cell background colour."""
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd {cell._element.nsmap["w"]} '
        f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:fill="{hex_colour}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def generate_memo(
    results: dict,
    output_path: str,
    firm_name: str = "Law Firm",
    transaction_type: str = "property",
    city: str = "Islamabad",
    document_names: list = None,
    flags: dict = None,
) -> str:
    """
    Generate a structured due diligence memo as a .docx file.
    Returns the path to the generated file.
    """
    doc = Document()
    document_names = document_names or []
    flags = flags or {}

    # ── Page margins ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Header ──────────────────────────────────────────────────────
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"{firm_name} — Legal Due Diligence System"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _style_run(header_para.runs[0], size=9, colour=COLOUR_ACCENT)

    # ── Cover section ───────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DUE DILIGENCE REVIEW MEMORANDUM")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOUR_PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(
        f"{transaction_type.title()} Transaction — {city.title()}"
    )
    run2.font.size = Pt(13)
    run2.font.color.rgb = COLOUR_ACCENT

    doc.add_paragraph()

    # ── Transaction summary table ───────────────────────────────────
    _add_heading(doc, "1. Transaction Summary", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    summary_data = [
        ("Prepared by",        firm_name),
        ("Date",               datetime.now().strftime("%d %B %Y")),
        ("Transaction type",   transaction_type.title()),
        ("City / Authority",   city.title()),
        ("Documents reviewed", str(len(document_names))),
        ("Questions assessed", str(results.get("total_questions", 0))),
    ]
    for i, (label, value) in enumerate(summary_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        _style_cell(row.cells[0], bold=True, colour=COLOUR_PRIMARY)
        _style_cell(row.cells[1])

    doc.add_paragraph()

    # ── Risk summary ────────────────────────────────────────────────
    _add_heading(doc, "2. Risk Summary", level=1)
    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = "Table Grid"
    risk_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["🔴 HIGH Risk", "🟡 MEDIUM Risk", "🟢 LOW Risk"]
    counts  = [
        results.get("high_risk_count",   0),
        results.get("medium_risk_count", 0),
        results.get("low_risk_count",    0),
    ]
    colours = ["AA2323", "B85C00", "1E6E37"]

    for i, (hdr, cnt, col) in enumerate(zip(headers, counts, colours)):
        cell = risk_table.rows[0].cells[i]
        cell.text = f"{hdr}\n{cnt} item(s)"
        _style_cell(cell, bold=True, colour=RGBColor.from_string(col))

    doc.add_paragraph()

    # ── Red flags ───────────────────────────────────────────────────
    red_flags = results.get("red_flags", [])
    _add_heading(doc, "3. Red Flags Detected", level=1)

    if not red_flags:
        p = doc.add_paragraph("No critical red flags detected.")
        p.runs[0].font.color.rgb = COLOUR_GREEN
    else:
        for flag in red_flags:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{flag['id']}] {flag['label']}")
            run.bold = True
            run.font.color.rgb = COLOUR_RED
            note = p.add_run(
                f"\n    Statute: {flag['statute']}"
                f"\n    Constitutional basis: {flag['article']}"
            )
            note.font.size = Pt(9)
            note.font.color.rgb = COLOUR_ACCENT

    doc.add_paragraph()

    # ── Clause-by-clause findings ───────────────────────────────────
    _add_heading(doc, "4. Clause-by-Clause Findings", level=1)

    findings = results.get("findings", [])
    for finding in findings:
        q_id   = finding.get("question_id", "")
        risk   = finding.get("risk_level", "LOW").upper()
        colour = _risk_colour(risk)

        # Question heading
        q_para = doc.add_paragraph()
        q_run  = q_para.add_run(
            f"Q{q_id}. {finding.get('question', '')[:120]}"
        )
        q_run.bold = True
        q_run.font.size = Pt(11)
        q_run.font.color.rgb = COLOUR_PRIMARY

        # Finding detail table
        detail = doc.add_table(rows=7, cols=2)
        detail.style = "Table Grid"

        rows_data = [
            ("Finding",               finding.get("finding",              "N/A")),
            ("Reasoning",             finding.get("reasoning",            "N/A")),
            ("Document citation",     finding.get("document_citation",    "N/A")),
            ("Statutory citation",    finding.get("statutory_citation",   "N/A")),
            ("Constitutional basis",  finding.get("constitutional_basis", "N/A")),
            ("Risk level",            risk),
            ("Recommendation",        finding.get("recommendation",       "N/A")),
        ]
        for i, (label, value) in enumerate(rows_data):
            row = detail.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            _style_cell(row.cells[0], bold=True, colour=COLOUR_PRIMARY)
            cell_colour = colour if label == "Risk level" else None
            _style_cell(row.cells[1], colour=cell_colour)

        doc.add_paragraph()

    # ── FBR / AML Compliance ────────────────────────────────────────
    _add_heading(doc, "5. FBR & AML Compliance Assessment", level=1)

    fbr_applicable  = flags.get("fbr_applicable",  False)
    aml_threshold   = flags.get("aml_threshold",   False)
    detected_value  = flags.get("detected_value_pkr", 0)
    aml_indicators  = flags.get("aml_indicators",  False)
    has_urdu        = flags.get("has_urdu",         False)
    has_ocr         = flags.get("has_ocr_pages",    False)

    fbr_table = doc.add_table(rows=6, cols=2)
    fbr_table.style = "Table Grid"

    def _yn(val: bool) -> str:
        return "YES — Action Required" if val else "No — Within Threshold"

    fbr_rows = [
        (
            "Detected transaction value",
            f"PKR {detected_value:,.0f}" if detected_value else "Not detected in documents",
        ),
        (
            "FBR withholding tax applicable\n(Section 236C/236K — above PKR 5M)",
            "YES — Withholding tax required (1% filers / 2% non-filers)"
            if fbr_applicable
            else "No — Below PKR 5,000,000 threshold",
        ),
        (
            "Enhanced AML due diligence required\n(AML Act 2010 — above PKR 10M)",
            "YES — Source of funds documentation required"
            if aml_threshold
            else "No — Below PKR 10,000,000 threshold",
        ),
        (
            "AML risk indicators detected",
            "YES — Review source of funds documentation"
            if aml_indicators
            else "None detected",
        ),
        (
            "Urdu content detected in documents",
            "YES — Multilingual processing applied"
            if has_urdu
            else "No — English-only document",
        ),
        (
            "OCR applied to scanned pages",
            "YES — Some pages were scanned and OCR-processed"
            if has_ocr
            else "No — All pages are text-based",
        ),
    ]

    for i, (label, value) in enumerate(fbr_rows):
        row = fbr_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        _style_cell(row.cells[0], bold=True, colour=COLOUR_PRIMARY)

        # Colour HIGH risk cells red
        needs_action = "YES" in value and "Action" in value or "required" in value.lower()
        _style_cell(
            row.cells[1],
            colour=COLOUR_RED if needs_action else None,
        )

    doc.add_paragraph()

    # Statutory note
    if fbr_applicable or aml_threshold:
        note = doc.add_paragraph()
        note_run = note.add_run(
            "Statutory Reference: Income Tax Ordinance 2001, Sections 236C "
            "and 236K (withholding tax on property transactions); "
            "Anti-Money Laundering Act 2010 (enhanced due diligence); "
            "SECP AML/CFT Regulations 2018 (designated non-financial businesses). "
            "Constitutional basis: Article 23 — Right to acquire property "
            "subject to applicable law."
        )
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = COLOUR_ACCENT
        note_run.italic = True
        doc.add_paragraph()

    # ── Missing documents ───────────────────────────────────────────
    _add_heading(doc, "6. Missing Documents", level=1)
    all_missing = []
    for f in findings:
        all_missing.extend(f.get("missing_documents", []))
    all_missing = list(set(all_missing))

    if not all_missing:
        doc.add_paragraph("No missing documents identified.")
    else:
        for doc_name in all_missing:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(doc_name).font.color.rgb = COLOUR_RED

    doc.add_paragraph()

    # ── Disclaimer ──────────────────────────────────────────────────
    _add_heading(doc, "7. Disclaimer", level=1)
    disclaimer = doc.add_paragraph(
        "This memorandum has been generated with the assistance of an "
        "AI-powered legal review system and must be reviewed, verified, "
        "and approved by a qualified Pakistani advocate before being relied "
        "upon. It does not constitute legal advice. All findings should be "
        "independently verified against original documents."
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.color.rgb = COLOUR_ACCENT

    # ── Save ────────────────────────────────────────────────────────
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Memo saved: {output_path}")
    return output_path


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = COLOUR_PRIMARY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)


def _style_run(run, size: int = 10, colour: RGBColor = None, bold: bool = False):
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    run.bold = bold


def _style_cell(cell, bold: bool = False, colour: RGBColor = None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
            run.bold = bold
            if colour:
                run.font.color.rgb = colour